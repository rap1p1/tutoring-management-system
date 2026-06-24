import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_cell_background(cell, fill_color="F2F2F2"):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color))
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Load existing template
doc = Document("Mau Cuon Bao Cao Do An CNPM.docx")

# Add Section
doc.add_heading("CHƯƠNG V. KIỂM THỬ HỆ THỐNG", level=1)
doc.add_paragraph("Mục này trình bày chi tiết các kịch bản kiểm thử (test cases) theo nghiệp vụ chính của hệ thống. Quá trình kiểm thử tập trung vào luồng phân công gia sư, báo vắng và tài chính.")

test_cases_to_include = [
    "TC-3.1", "TC-3.4", "TC-4.2", "TC-4.6", "TC-5.2", "TC-5.4", "TC-6.1"
]

with open("test_cases.md", "r", encoding="utf-8") as f:
    lines = f.readlines()

current_table = []
in_table = False
capture_next = False
current_heading = ""
include_this_tc = False

for line in lines:
    line = line.strip()
    
    if line.startswith("## "):
        # Only add relevant groups
        if "LUỒNG NGHIỆP VỤ" in line or "BÁO DẠY" in line or "BÁO NGHỈ" in line or "TÀI CHÍNH" in line:
            current_heading = line.replace("## ", "")
            capture_next = True
        else:
            capture_next = False

    if not capture_next:
        continue

    if line.startswith("## "):
        doc.add_heading(current_heading, level=2)
    elif line.startswith("### "):
        tc_id = line.split(":")[0].replace("### ", "").strip()
        if tc_id in test_cases_to_include:
            doc.add_heading(line.replace("### ", ""), level=3)
            include_this_tc = True
        else:
            include_this_tc = False
    elif include_this_tc and line.startswith("|") and not line.startswith("|---"):
        if not in_table:
            in_table = True
            current_table = []
        cells = [c.strip() for c in line.split("|")[1:-1]]
        current_table.append(cells)
    elif include_this_tc and line.startswith("|---"):
        pass
    elif include_this_tc:
        if in_table:
            table = doc.add_table(rows=len(current_table), cols=len(current_table[0]))
            for i, row in enumerate(current_table):
                for j, cell_text in enumerate(row):
                    cell = table.cell(i, j)
                    cell.text = cell_text.replace("`", "") # Remove markdown backticks
                    if i == 0:
                        set_cell_background(cell, "D9D9D9") # Header gray
                    else:
                        set_cell_background(cell, "F2F2F2") # Row gray
            in_table = False
            current_table = []
            doc.add_paragraph("") # Space
            
        if line and not line.startswith(">"):
            if "Mong đợi" in line or "✅" in line:
                p = doc.add_paragraph()
                r = p.add_run(line.replace("✅", "").replace("**", ""))
                r.bold = True
            else:
                doc.add_paragraph(line.replace("**", ""))

# Add the priority matching test case custom logic
doc.add_heading("TC-3.6: Kiểm tra độ ưu tiên khi chọn gia sư để ghép lớp (Custom Test)", level=3)
table = doc.add_table(rows=3, cols=3)
headers = ["Bước", "Thao tác", "Kết quả mong đợi"]
for j, h in enumerate(headers):
    table.cell(0, j).text = h
    set_cell_background(table.cell(0, j), "D9D9D9")

table.cell(1, 0).text = "1"
table.cell(1, 1).text = "Khởi tạo 2 tài khoản Gia sư: GS1 trùng 2 lịch rảnh (có đăng ký dạy môn Toán), GS2 trùng 0 lịch rảnh (nhưng không đăng ký dạy môn Toán)"
table.cell(1, 2).text = "Tạo và phê duyệt thành công"
set_cell_background(table.cell(1, 0), "F2F2F2")
set_cell_background(table.cell(1, 1), "F2F2F2")
set_cell_background(table.cell(1, 2), "F2F2F2")

table.cell(2, 0).text = "2"
table.cell(2, 1).text = "Sử dụng tài khoản học viên gửi yêu cầu ghép lớp môn Toán. Nhân viên QL bấm 'Ghép Lớp' để kiểm tra danh sách gợi ý."
table.cell(2, 2).text = "Hệ thống hiển thị danh sách GS gợi ý được sắp xếp ưu tiên dựa trên môn đăng ký và mức độ trùng lịch."
set_cell_background(table.cell(2, 0), "F2F2F2")
set_cell_background(table.cell(2, 1), "F2F2F2")
set_cell_background(table.cell(2, 2), "F2F2F2")
doc.add_paragraph("")

p = doc.add_paragraph()
r = p.add_run("Kết quả thực tế (Qua quá trình kiểm thử tự động): ")
r.bold = True
p.add_run("Hệ thống hiển thị GS1 lên đầu danh sách vì trùng 2 lịch rảnh (matchCount = 2). Tuy nhiên, nếu đổi kịch bản: GS2 trùng 3 lịch rảnh nhưng KHÔNG đăng ký dạy Toán, hệ thống vẫn ưu tiên GS2 lên đầu tiên. Thuật toán gợi ý hiện tại ở frontend (dựa vào hàm getMatchCount) CHƯA có logic kiểm tra gia sư có đăng ký dạy môn học đó hay không.")

p_err = doc.add_paragraph()
r_err = p_err.add_run("=> Đánh giá: Có lỗi logic trong việc gợi ý gia sư (cần phải filter môn học trước khi tính điểm trùng lịch rảnh).")
r_err.font.color.rgb = RGBColor(255, 0, 0)
r_err.bold = True

doc.add_paragraph("")
doc.add_paragraph("[Chèn ảnh giao diện ghép lớp có dropdown ở đây]")
p_img = doc.add_paragraph()
r_img = p_img.add_run(": bảng 5.1 Giao diện ghép lớp và gợi ý gia sư")
r_img.italic = True

doc.save("Mau_Cuon_Bao_Cao_Do_An_CNPM_Updated.docx")
print("Report generated: Mau_Cuon_Bao_Cao_Do_An_CNPM_Updated.docx")
